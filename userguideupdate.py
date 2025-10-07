import os
import re
import shutil
import win32com.client
from docx import Document
from docx.shared import Inches, Pt  # Import for font sizing

def clear_com_cache():
    cache_dir = win32com.client.gencache.GetGeneratePath()
    if os.path.exists(cache_dir):
        print(f"Clearing COM cache folder: {cache_dir}")
        shutil.rmtree(cache_dir)

# Call this once, early in your program.
clear_com_cache()

def get_run_font_properties(run):
    if run and run.font:
        return {
            'name': run.font.name,
            'size': run.font.size,
            'bold': run.font.bold,
            'italic': run.font.italic
        }
    return {}

def apply_font_properties(run, properties):
    if properties:
        if properties.get('name') is not None:
            run.font.name = properties['name']
        if properties.get('size') is not None:
            run.font.size = properties['size']
        if properties.get('bold') is not None:
            run.font.bold = properties['bold']
        if properties.get('italic') is not None:
            run.font.italic = properties['italic']

def convert_rtf_to_docx(rtf_path, docx_path):
    word = None # Initialize word to None for proper cleanup
    try:
        word = win32com.client.gencache.EnsureDispatch("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(rtf_path)
        # 12 = wdFormatDocumentDefault (DOCX)
        doc.SaveAs(docx_path, FileFormat=12)
        doc.Close()
        print(f"Successfully converted RTF to DOCX: {rtf_path} -> {docx_path}")
    except Exception as e:
        print(f"Error converting RTF to DOCX: {e}")
    finally:
        if word:
            word.Quit()

def convert_docx_to_rtf(docx_path, rtf_output_path):
    word = None  # Initialize word to None for proper cleanup
    try:
        word = win32com.client.gencache.EnsureDispatch("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(docx_path)
        doc.SaveAs(rtf_output_path, FileFormat=6)  # 6 = wdFormatRTF
        doc.Close()
        print(f"Successfully converted DOCX to RTF: {docx_path} -> {rtf_output_path}")
    except Exception as e:
        print(f"Error converting DOCX to RTF: {e}")
    finally:
        if word:
            word.Quit()

def modify_docx_title(docx_path, new_title):
    try:
        doc = Document(docx_path)
        target_prefix = "Forms on Download for"
        found_and_replaced = False

        for para in doc.paragraphs:
            if para.text.strip().startswith(target_prefix):
                # Preserve existing formatting from the first run if available
                first_run_props = get_run_font_properties(para.runs[0] if para.runs else None)
                para_style = para.style

                # Clear existing content and add the new content as a single run
                # This ensures consistent formatting for the new title phrase
                para.clear()
                run = para.add_run(f"{target_prefix} {new_title}")
                apply_font_properties(run, first_run_props) # Apply preserved font properties
                para.style = para_style # Reapply paragraph style

                print(f"Updated title paragraph: {para.text}")
                found_and_replaced = True
                break # Only update the first matching paragraph

        if not found_and_replaced:
            print(f"Warning: Target prefix '{target_prefix}' not found in {docx_path} for title modification.")

        doc.save(docx_path)
    except Exception as e:
        print(f"Error modifying DOCX title: {e}")

def modify_docx_relnotes(docx_path, new_release_notes):
    try:
        doc = Document(docx_path)
        # Define the target text pattern.
        target_pattern = r"Changes to Forms on Download Release XX, Month Year"

        found_and_replaced = False

        for para in doc.paragraphs:
            original_text = para.text  # store the original paragraph text
            match = re.search(target_pattern, original_text)
            if match:
                # Prepare to build the new content while preserving formatting.
                new_paragraph_runs_content = []
                current_char_index = 0

                for run in para.runs:
                    # Find the starting index of the matched pattern from the current index.
                    start_match_index = original_text.find(match.group(0), current_char_index)
                    if start_match_index != -1 and current_char_index <= start_match_index:
                        # Add any text that appears before the match.
                        if current_char_index < start_match_index:
                            new_paragraph_runs_content.append({
                                'text': original_text[current_char_index:start_match_index],
                                'props': get_run_font_properties(run)
                            })

                        # Replace the matched text with new_release_notes.
                        new_paragraph_runs_content.append({
                            'text': new_release_notes,
                            'props': get_run_font_properties(run)
                        })

                        # Update the current index past the replaced part.
                        current_char_index = start_match_index + len(match.group(0))

                    # Append any remaining text of the run if not already handled.
                    if current_char_index < len(original_text):
                        # Determine the next occurrence of the target pattern.
                        next_match_index = original_text.find(match.group(0), current_char_index)
                        end_index = next_match_index if next_match_index != -1 else len(original_text)
                        if current_char_index < end_index:
                            new_paragraph_runs_content.append({
                                'text': original_text[current_char_index:end_index],
                                'props': get_run_font_properties(run)
                            })
                            current_char_index = end_index

                # Save the current style and clear the paragraph's content.
                para_style = para.style
                para.clear()  # remove all runs from this paragraph

                # Rebuild the paragraph with the new runs and apply the preserved formatting.
                for item in new_paragraph_runs_content:
                    new_run = para.add_run(item['text'])
                    apply_font_properties(new_run, item['props'])
                para.style = para_style

                print(f"Updated in-paragraph notes: {para.text}")
                found_and_replaced = True
                break  # update only the first matching paragraph

        if not found_and_replaced:
            print(f"Warning: Target pattern for in-paragraph title not found or matched in {docx_path}.")

        doc.save(docx_path)
    except Exception as e:
        print(f"Error modifying DOCX in-paragraph title: {e}")

def modify_docx_title_inpara(docx_path, new_title):
    try:
        doc = Document(docx_path)
        # Using a more robust regex to capture the specific part to replace
        # This assumes "Forms on Download for" is followed by a title that ends before " is designed"
        target_pattern = r"Forms on Download for (.*?) is designed "
        replacement_start = "Forms on Download for "
        replacement_end = " is designed "
        found_and_replaced = False

        for para in doc.paragraphs:
            original_text = para.text # Keep original text to check if it matches the full phrase
            match = re.search(target_pattern, original_text)
            if match:
                # Capture existing formatting of the paragraph's runs
                # and reconstruct the paragraph while replacing the specific part
                new_paragraph_runs_content = []
                current_char_index = 0
                for run in para.runs:
                    run_text_lower = run.text.lower() # Convert to lower for case-insensitive matching if needed
                    # If this run contains part of the matched pattern, handle replacement
                    start_match_index = original_text.find(match.group(0), current_char_index)
                    if start_match_index != -1 and current_char_index <= start_match_index:
                        # Add text before the match
                        if current_char_index < start_match_index:
                            new_paragraph_runs_content.append({'text': original_text[current_char_index:start_match_index], 'props': get_run_font_properties(run)})
                        
                        # Add the new title part
                        new_paragraph_runs_content.append({'text': f"{replacement_start}{new_title}{replacement_end}", 'props': get_run_font_properties(run)})
                        
                        current_char_index = start_match_index + len(match.group(0))
                    
                    # Add remaining text from the run if not part of the replaced section
                    if current_char_index < len(original_text) and current_char_index < (original_text.find(match.group(0), current_char_index) if original_text.find(match.group(0), current_char_index) != -1 else len(original_text)):
                         new_paragraph_runs_content.append({'text': original_text[current_char_index: len(original_text)], 'props': get_run_font_properties(run)})
                         current_char_index = len(original_text)


                # Rebuild the paragraph with new content and preserved formatting
                para_style = para.style
                para.clear()
                for item in new_paragraph_runs_content:
                    new_run = para.add_run(item['text'])
                    apply_font_properties(new_run, item['props'])
                para.style = para_style

                print(f"Updated in-paragraph title: {para.text}")
                found_and_replaced = True
                break  # Only update the first matching paragraph

        if not found_and_replaced:
            print(f"Warning: Target pattern for in-paragraph title not found or matched in {docx_path}.")

        doc.save(docx_path)
    except Exception as e:
        print(f"Error modifying DOCX in-paragraph title: {e}")


def modify_docx_sample(docx_path, new_sample):
    try:
        doc = Document(docx_path)
        target_prefix = "Each file contains a single form or checklist. Each file name corresponds directly to the section number of the form or checklist. For example, "
        target_suffix = " Forms on Disc, you may refer to the corresponding section in the publication for specific information about each form or checklist. A complete list of the forms and checklists is included with Forms on Disc in FORMLIST.rtf."
        found_and_replaced = False

        for para in doc.paragraphs:
            if para.text.strip().startswith(target_prefix):
                first_run_props = get_run_font_properties(para.runs[0] if para.runs else None)
                para_style = para.style

                # Construct and modify the full paragraph text
                full_text = f"{target_prefix}{new_sample}{target_suffix}"
                full_text = full_text.replace("Disc", "Download").replace("disc", "download")

                para.clear()
                run = para.add_run(full_text)
                apply_font_properties(run, first_run_props)
                para.style = para_style

                print(f"Updated sample paragraph: {para.text}")
                found_and_replaced = True
                break

        if not found_and_replaced:
            print(f"Warning: Target prefix '{target_prefix}' not found in {docx_path} for sample modification.")
        doc.save(docx_path)
    except Exception as e:
        print(f"Error modifying DOCX sample: {e}")

# Consolidated function for modifying specific content (Release, Revised, New, Deleted Forms)
# This pattern is repeated, so a single helper function makes sense.
def _modify_docx_section(docx_path, target_placeholder, new_content, force_non_bold_new_content=False):
    try:
        doc = Document(docx_path)
        found_and_replaced = False

        for para in doc.paragraphs:
            if target_placeholder in para.text:
                para_style = para.style
                first_run_props = get_run_font_properties(para.runs[0] if para.runs else None)

                para.clear() # Clear paragraph content.

                # Add run for target placeholder with original formatting.
                placeholder_run = para.add_run(target_placeholder)
                apply_font_properties(placeholder_run, first_run_props)

                para.add_run("\n") # Add a newline.

                # Add run for new content.
                content_run = para.add_run(new_content)
                apply_font_properties(content_run, first_run_props) # Apply original properties first
                if force_non_bold_new_content:
                    content_run.font.bold = False # Override to force non-bold

                para.style = para_style
                print(f"Updated section for '{target_placeholder}': {para.text[:50]}...") # Print first 50 chars for brevity
                found_and_replaced = True
                break # Update only the first matching paragraph

        if not found_and_replaced:
            print(f"Warning: Target placeholder '{target_placeholder}' not found in {docx_path}.")

        doc.save(docx_path)
    except Exception as e:
        print(f"Error modifying DOCX for '{target_placeholder}' section: {e}")

# Step 2b: Modify release notes while preserving paragraph formatting
# def modify_docx_relnotes(docx_path, new_release_notes):
#     _modify_docx_section(docx_path, "Changes to Forms on Download Release XX, Month Year", new_release_notes)

# Step 2c: Modify revised forms notes (forcing non-bold for new content)
def modify_docx_revisenotes(docx_path, new_revised_forms_content):
    _modify_docx_section(docx_path, "REVISED:", new_revised_forms_content, force_non_bold_new_content=True)

# Step 2d: Modify new forms notes
def modify_docx_newforms(docx_path, new_forms_content):
    _modify_docx_section(docx_path, "NEW:", new_forms_content, force_non_bold_new_content=True)

# Step 2e: Modify deleted forms notes
def modify_docx_deletednotes(docx_path, new_deleted_forms_content):
    _modify_docx_section(docx_path, "DELETED:", new_deleted_forms_content, force_non_bold_new_content=True)

def userguide_update(self):
    temp_entry = self.entry1() if callable(self.entry1) else self.entry1
    folder_entry = temp_entry.get() if hasattr(temp_entry, "get") else temp_entry

    # initial_dir = r"C:\Users\LABRADBM\Downloads\Local\YB\Python\FOD\I drive"
    initial_dir = r"\\fabwebd5.net\neptune\DataConversion\Prod\Secondary"
    # base_dir = os.path.join(initial_dir, folder_entry)
    base_dir = os.path.join(initial_dir, folder_entry, "FOD")
    report_directory = os.path.join(base_dir, 'Report')
    os.makedirs(report_directory, exist_ok=True)

    # Define file paths.
    rtf_path = os.path.join(report_directory, 'ref_userguide.rtf')
    docx_path = os.path.join(report_directory, 'temp.docx')
    modified_rtf_path = os.path.join(report_directory, 'modified_userguide.rtf')
    pub_title_path = os.path.join(report_directory, 'PubTitle.txt')
    pub_sample_path = os.path.join(report_directory, 'PubSample.txt')
    release_notes_path = os.path.join(report_directory, 'ReleaseNotes.txt')
    collected_forms_path = os.path.join(report_directory, 'CollectedForms.txt')

    # Initialize variables to 'Error' state
    new_title = "Error"
    new_sample = "Error"
    new_release_notes = "Error"
    new_revised_forms_content = "Error"
    new_forms_content = "Error"
    new_deleted_forms_content = "Error"

    try:
        # Read all necessary content from text files
        # Read PubTitle.txt
        if os.path.exists(pub_title_path):
            with open(pub_title_path, 'r', encoding='utf-8') as f:
                pub_title_content = f.read()
                match_title = re.search(r"Extracted Title: (.+)", pub_title_content)
                new_title = match_title.group(1).strip() if match_title else "Error: Title not found"
        else:
            print(f"Warning: {pub_title_path} not found.")

        # Read PubSample.txt
        if os.path.exists(pub_sample_path):
            with open(pub_sample_path, 'r', encoding='utf-8') as f:
                pub_sample_content = f.read()
                match_sample = re.search(r"Extracted Sample: (.+)", pub_sample_content)
                new_sample = match_sample.group(1).strip() if match_sample else "Error: Sample not found"
        else:
            print(f"Warning: {pub_sample_path} not found.")

        # Read ReleaseNotes.txt
        if os.path.exists(release_notes_path):
            with open(release_notes_path, 'r', encoding='utf-8') as f:
                release_notes_content = f.read()
                # Assuming Release Notes: is the first line, then the content
                match_release = re.search(r"Release Notes:\n(.+)", release_notes_content, re.DOTALL)
                new_release_notes = match_release.group(1).strip() if match_release else "Error: Release notes not found"
        else:
            print(f"Warning: {release_notes_path} not found.")

        # Read CollectedForms.txt once
        if os.path.exists(collected_forms_path):
            with open(collected_forms_path, 'r', encoding='utf-8') as f:
                collected_forms_content = f.read()

            # Capture content between "Revised Forms:" and "New Forms:" (non-greedy)
            match_revised = re.search(r"Revised Forms:\n(.*?)\nNew Forms:", collected_forms_content, re.DOTALL)
            new_revised_forms_content = match_revised.group(1).strip() if match_revised else "Error: Revised forms not found"

            # Capture content between "New Forms:" and "Deleted Forms:"
            match_new = re.search(r"New Forms:\n(.*?)\nDeleted Forms:", collected_forms_content, re.DOTALL)
            new_forms_content = match_new.group(1).strip() if match_new else "Error: New forms not found"

            # Capture all text after "Deleted Forms:"
            match_deleted = re.search(r"Deleted Forms:\n(.+)", collected_forms_content, re.DOTALL)
            new_deleted_forms_content = match_deleted.group(1).strip() if match_deleted else "Error: Deleted forms not found"
        else:
            print(f"Warning: {collected_forms_path} not found.")


        print(f"New Title: '{new_title}'")
        print(f"New Sample: '{new_sample}'")
        print(f"New Release Notes: '{new_release_notes}'")
        print(f"New Revised Forms: '{new_revised_forms_content}'")
        print(f"New Forms: '{new_forms_content}'")
        print(f"New Deleted Forms: '{new_deleted_forms_content}'")


        # Convert RTF to DOCX, perform modifications, then convert back to RTF.
        if not os.path.exists(rtf_path):
            print(f"Error: Source RTF file not found at {rtf_path}")
            return # Exit if the source RTF is missing

        convert_rtf_to_docx(rtf_path, docx_path)

        # Only proceed with DOCX modifications if conversion was successful
        if os.path.exists(docx_path):
            modify_docx_title(docx_path, new_title)
            # The logic for modify_docx_title_inpara might need further refinement
            # based on actual document structure for robust replacements.
            modify_docx_title_inpara(docx_path, new_title)
            modify_docx_sample(docx_path, new_sample)
            modify_docx_relnotes(docx_path, new_release_notes)
            modify_docx_revisenotes(docx_path, new_revised_forms_content)
            modify_docx_newforms(docx_path, new_forms_content)
            modify_docx_deletednotes(docx_path, new_deleted_forms_content)

            convert_docx_to_rtf(docx_path, modified_rtf_path)
            print(f"✅ Modified RTF saved at: {modified_rtf_path}")
        else:
            print(f"❌ DOCX conversion failed, skipping document modifications.")

    except Exception as e:
        print(f"❌ Main processing error in userguide_update: {e}")
    finally:
        # Clean up temporary DOCX file if it exists
        if os.path.exists(docx_path):
            try:
                os.remove(docx_path)
                print(f"Cleaned up temporary DOCX file: {docx_path}")
            except Exception as e:
                print(f"Error cleaning up temporary DOCX file {docx_path}: {e}")

